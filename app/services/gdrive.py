import os
import re
import io
import base64
import httpx
import logging
from typing import List, Dict, Tuple

logger = logging.getLogger(__name__)

DRIVE_API    = "https://www.googleapis.com/drive/v3"
EXPORT_API   = "https://www.googleapis.com/drive/v3/files/{id}/export"
DOWNLOAD_API = "https://www.googleapis.com/drive/v3/files/{id}?alt=media"

SUPPORTED_MIME = {
    "application/pdf":                                                          "pdf",
    "text/plain":                                                               "text",
    "text/markdown":                                                            "text",
    "application/vnd.google-apps.document":                                    "gdoc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


def extract_folder_id(url: str):
    patterns = [
        r"drive\.google\.com/drive/folders/([a-zA-Z0-9_-]+)",
        r"drive\.google\.com/open\?id=([a-zA-Z0-9_-]+)",
        r"id=([a-zA-Z0-9_-]+)",
    ]
    for p in patterns:
        m = re.search(p, url)
        if m:
            return m.group(1)
    return None


async def list_folder_files(folder_id: str, api_key: str) -> List[Dict]:
    files = []
    page_token = None

    async with httpx.AsyncClient(timeout=30) as client:
        while True:
            params = {
                "q": f"'{folder_id}' in parents and trashed=false",
                "fields": "nextPageToken, files(id, name, mimeType, size)",
                "key": api_key,
                "pageSize": 100,
            }
            if page_token:
                params["pageToken"] = page_token

            resp = await client.get(f"{DRIVE_API}/files", params=params)

            if resp.status_code == 403:
                raise Exception("Folder is not public. Set sharing to 'Anyone with the link can view'.")
            if not resp.is_success:
                raise Exception(f"Drive API error {resp.status_code}: {resp.text}")

            data = resp.json()
            for f in data.get("files", []):
                if f.get("mimeType") in SUPPORTED_MIME:
                    files.append(f)
                else:
                    logger.info(f"Skipping: {f['name']} ({f.get('mimeType')})")

            page_token = data.get("nextPageToken")
            if not page_token:
                break

    return files


def _extract_docx_text(content: bytes) -> str:
    """Extract plain text from .docx bytes using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n\n".join(parts)
        logger.info(f"Extracted {len(text)} chars from docx")
        return text
    except Exception as e:
        logger.error(f"docx extraction failed: {e}")
        raise Exception(f"Could not parse Word document: {e}")


async def download_file(file_id: str, mime_type: str, api_key: str) -> Tuple[str, str]:
    """
    Downloads a file and returns (content, file_type).
    file_type is one of: 'pdf' | 'text'
    For PDFs: content is base64-encoded bytes
    For everything else: content is plain text
    """
    async with httpx.AsyncClient(timeout=60, follow_redirects=True) as client:

        if mime_type == "application/vnd.google-apps.document":
            # Export Google Doc as plain text directly
            resp = await client.get(
                EXPORT_API.format(id=file_id),
                params={"mimeType": "text/plain", "key": api_key}
            )
            if not resp.is_success:
                raise Exception(f"Export failed {resp.status_code}: {resp.text[:200]}")
            text = resp.content.decode("utf-8", errors="replace").strip()
            logger.info(f"Google Doc exported: {len(text)} chars")
            return text, "text"

        else:
            # Direct binary download for PDF, docx, txt
            resp = await client.get(
                DOWNLOAD_API.format(id=file_id),
                params={"key": api_key}
            )
            if not resp.is_success:
                raise Exception(f"Download failed {resp.status_code}: {resp.text[:200]}")

            if mime_type == "application/pdf":
                # Return base64 for pdfminer processing
                b64 = base64.b64encode(resp.content).decode()
                logger.info(f"PDF downloaded: {len(resp.content)} bytes")
                return b64, "pdf"

            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Parse docx binary → plain text
                text = _extract_docx_text(resp.content)
                return text, "text"

            else:
                # Plain text files
                text = resp.content.decode("utf-8", errors="replace").strip()
                logger.info(f"Text file downloaded: {len(text)} chars")
                return text, "text"